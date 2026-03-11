import os
import docx
import json
import requests
import re
import time
from concurrent.futures import ThreadPoolExecutor
from tqdm import tqdm

# Settings
WORD_DIR = "word_planlar"
JSON_DIR = "json_planlar"
MODEL = "llama3.2"
OLLAMA_URL = "http://localhost:11434/api/generate"
MAX_WORKERS = 1  # Sequential processing is safer for slow CPU/GPU
TIMEOUT = 900    # 15 minutes for very slow inference

os.makedirs(JSON_DIR, exist_ok=True)

def print_flush(msg):
    print(msg, flush=True)

def extract_relevant_text(doc_path):
    try:
        doc = docx.Document(doc_path)
        metin_parcalari = []
        
        # İlk paragraflar
        for p in doc.paragraphs[:10]:
            text = p.text.strip()
            if text and not (text.count('.') > 10 or "Müdürü" in text or "Öğretmeni" in text):
                metin_parcalari.append(text)
        
        # Tabloları tara
        for tablo in doc.tables:
            for satir in tablo.rows:
                hucre_metinleri = [hucre.text.strip() for hucre in satir.cells if hucre.text.strip()]
                # İmza alanlarını ve çok sayıda nokta içeren satırları atla
                line = " | ".join(hucre_metinleri)
                if len(hucre_metinleri) > 1 and not (line.count('.') > 10 or "UYGUNDUR" in line):
                    metin_parcalari.append(line)
        
        # Metni 3000 karakterle sınırla (Inferansı hızlandırmak için)
        return "\n".join(metin_parcalari)[:3000]
    except Exception as e:
        return f"Error reading file: {e}"

def process_file(dosya_adi):
    if not dosya_adi.endswith(".docx"):
        return None
    
    input_path = os.path.join(WORD_DIR, dosya_adi)
    
    # Text extraction
    tam_metin = extract_relevant_text(input_path)
    if "Error reading file" in tam_metin or len(tam_metin) < 100:
        return f"Skipped (Too short or error): {dosya_adi}"

    prompt = f"""Aşağıdaki metin bir öğretmenin yıllık ders planıdır. Lütfen metni analiz et ve aşağıdaki JSON formatında çıktı ver.
    'ders' alanını küçük harflerle, Türkçe karakter kullanmadan ve boşluk yerine alt tire ile yaz.
    {{
      "sinif": "7. Sınıf",
      "ders": "ingilizce",
      "plan": [ 
        {{
          "hafta": 1, 
          "tarih": "Eylül 1-5", 
          "kazanim": "Öğrenci selamlaşma ifadelerini kullanır", 
          "kod": "E7.1.L1"
        }} 
      ]
    }}
    Metin: {tam_metin}"""

    try:
        response = requests.post(OLLAMA_URL, json={
            "model": MODEL,
            "prompt": prompt,
            "format": "json", # Ollama JSON mode
            "stream": False
        }, timeout=TIMEOUT)
        
        if response.status_code == 200:
            veri = response.json().get("response", "{}")
            # Ollama string döndürür, parse etmemiz lazım
            try:
                json_data = json.loads(veri)
                # Sınıf ve ders bilgisini temizle (Türkçe karakterleri kaldır, küçük harf yap)
                def clean_name(text):
                    tr_map = str.maketrans("çğıöşü ", "cgiosu_")
                    return str(text).lower().translate(tr_map).replace(".", "")

                sinif = clean_name(json_data.get('sinif', 'unknown'))
                ders = clean_name(json_data.get('ders', 'unknown'))
                
                output_filename = f"grade_{sinif}_{ders}.json"
                output_path = os.path.join(JSON_DIR, output_filename)
                
                with open(output_path, "w", encoding="utf-8") as f:
                    json.dump(json_data, f, ensure_ascii=False, indent=2)
                return f"Success: {dosya_adi} -> {output_filename}"
            except json.JSONDecodeError:
                return f"Error (Invalid JSON): {dosya_adi}"
        else:
            return f"Error (HTTP {response.status_code}): {dosya_adi}"
            
    except Exception as e:
        return f"Error ({dosya_adi}): {e}"

def main():
    # Test için sadece plan_100 dosyasını seç
    test_files = ["plan_100.docx"]
    dosyalar = [f for f in test_files if os.path.exists(os.path.join(WORD_DIR, f))]
    toplam = len(dosyalar)
    print_flush(f"Tek dosya test modu: {toplam} dosya işleniyor... ({', '.join(dosyalar)})\n")
    
    with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
        for i, result in enumerate(executor.map(process_file, dosyalar), 1):
            if result:
                print_flush(f"[{i}/{toplam}] {result}")

if __name__ == "__main__":
    main()
