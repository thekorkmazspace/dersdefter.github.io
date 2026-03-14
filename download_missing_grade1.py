#!/usr/bin/env python3
import requests
import docx
import json
import time
import os
import re
from io import BytesIO

os.makedirs("json_planlar", exist_ok=True)

# Grade 1 missing IDs
missing_ids = [1949, 1428, 1539, 1622, 1697, 1669, 1953]

# Known ders names for each ID (for fallback)
id_to_ders = {
    1949: "Beden Eğitimi ve Oyun (MEB)",
    1428: "Müzik (Maarif)",
    1539: "Serbest Etkinlikler",
    1622: "Matematik (MEB)",
    1697: "Türkçe (MEB)",
    1669: "Hayat Bilgisi (MEB)",
    1953: "Görsel Sanatlar (MEB)",
}


def slugify(s):
    s = s.lower()
    s = s.replace("ı", "i").replace("ü", "u").replace("ö", "o")
    s = s.replace("ş", "s").replace("ç", "c").replace("ğ", "g")
    s = re.sub(r"[^a-z0-9]", "_", s)
    return re.sub(r"_+", "_", s).strip("_")


def extract_title_info(doc):
    """Extract sinif and ders from the first table header row."""
    sinif = "1"
    ders = ""
    if doc.tables:
        t = doc.tables[0]
        if t.rows:
            header_text = t.rows[0].cells[0].text.strip()
            # Try to extract sinif
            sinif_match = re.search(r"(\d+)\.\s*SINIF", header_text, re.IGNORECASE)
            if sinif_match:
                sinif = sinif_match.group(1)
            # Try to extract ders name - it's usually after "SINIF" and before "DERS"
            ders_match = re.search(
                r"\d+\.\s*SINIF\s+(.+?)\s+DERS", header_text, re.IGNORECASE
            )
            if ders_match:
                ders = ders_match.group(1).strip()
    return sinif, ders


def parse_table(doc, plan_id):
    """Parse the main plan table from the DOCX document."""
    sinif, ders = extract_title_info(doc)

    if not ders:
        ders = id_to_ders.get(plan_id, "Bilinmeyen Ders")

    plan = []

    if not doc.tables:
        return sinif, ders, plan

    t = doc.tables[0]
    rows = t.rows

    # Find the data start row (skip header rows - usually rows 0,1,2)
    # Row 2 typically has column names: AY, HAFTA, SAAT, ...
    data_start = 3
    for i, row in enumerate(rows[:5]):
        cells = [c.text.strip() for c in row.cells]
        if cells and cells[0] in ("AY", "ay"):
            data_start = i + 1
            break

    # Determine column count
    if not rows:
        return sinif, ders, plan

    num_cols = len(rows[0].cells)

    # Column mapping based on observed structure (11 columns):
    # 0=AY, 1=HAFTA, 2=SAAT, 3=ÜNİTE/TEMA, 4=KONU, 5=KAZANIM, 6=SÜREÇ BİLEŞENLERİ,
    # 7=SDB, 8=OB, 9=DEĞERLER, 10=BELİRLİ GÜN
    # Some docs may have 7 columns:
    # 0=AY, 1=HAFTA, 2=SAAT, 3=ÜNİTE, 4=KONU, 5=KAZANIM, 6=BELİRLİ GÜN

    for row in rows[data_start:]:
        cells = [c.text.strip() for c in row.cells]
        if not cells:
            continue

        ay = cells[0] if len(cells) > 0 else ""
        hafta = cells[1] if len(cells) > 1 else ""
        saat = cells[2] if len(cells) > 2 else ""
        unite = cells[3] if len(cells) > 3 else ""
        konu = cells[4] if len(cells) > 4 else ""
        kazanim = cells[5] if len(cells) > 5 else ""

        # BELİRLİ GÜN is last column
        belirli_gun = cells[-1] if len(cells) > 6 else ""

        # Skip completely empty rows
        if not ay and not hafta and not unite and not kazanim:
            continue

        # Skip rows that are just repeated header rows
        if ay.upper() in ("AY", "SÜRE") or hafta.upper() in ("HAFTA", "SÜRE"):
            continue

        entry = {
            "AY": ay,
            "HAFTA": hafta,
            "SAAT": saat,
            "ÜNİTE": unite,
            "KONU": konu,
            "KAZANIM": kazanim,
        }
        if belirli_gun:
            entry["BELİRLİ GÜN VE HAFTALAR"] = belirli_gun

        plan.append(entry)

    return sinif, ders, plan


for plan_id in missing_ids:
    url = f"https://defterdoldur.com/planword/{plan_id}"
    print(f"[{plan_id}] İndiriliyor: {url}")

    try:
        response = requests.get(url, timeout=15)
        if response.status_code != 200:
            print(f"[{plan_id}] HTTP {response.status_code} - Atlanıyor")
            continue

        try:
            doc = docx.Document(BytesIO(response.content))
        except Exception as e:
            print(f"[{plan_id}] DOCX okuma hatası: {e}")
            continue

        sinif, ders, plan = parse_table(doc, plan_id)

        if not ders:
            ders = id_to_ders.get(plan_id, f"plan_{plan_id}")

        veri = {
            "sinif": sinif,
            "ders": ders,
            "plan": plan,
        }

        ders_slug = slugify(ders)
        dosya_adi = f"grade_{sinif}_{ders_slug}_{plan_id}.json"
        dosya_yolu = os.path.join("json_planlar", dosya_adi)

        with open(dosya_yolu, "w", encoding="utf-8") as f:
            json.dump(veri, f, ensure_ascii=False, indent=2)

        print(f"[{plan_id}] Kaydedildi: {dosya_adi} ({len(plan)} satır)")

    except Exception as e:
        print(f"[{plan_id}] Hata: {e}")

    time.sleep(1)

print("\nTamamlandı!")
